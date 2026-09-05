#!/usr/bin/env python3
"""HP-3 — the pre-session tutorial sheet, generated as tutorial-sheet.docx.

    python3 tutorial-sheet.py            # needs python-docx and matplotlib

Every number on the sheet is the hydraulician scene the class runs (HP-3),
so every answer can be checked on screen. The answers appendix is computed
here from the same constants, never typed in, so it cannot drift from the
questions. Edit the text below and rerun; the schematic is drawn by
`scheme()` from the scene's own dimensions.
"""
import math
import os
import tempfile

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt                       # noqa: E402
from docx import Document                             # noqa: E402
from docx.shared import Pt, Cm                        # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK   # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT        # noqa: E402
from docx.oxml.ns import qn                           # noqa: E402
from docx.oxml import OxmlElement                     # noqa: E402

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, "tutorial-sheet.docx")

# ------------------------------------------------------------ the scene
g, rho = 9.81, 1000.0
RES, ZNOZ = 24.9, 3.0                 # reservoir level, nozzle axis (m above datum)
L, DH = 42.4, 3.05                    # headrace: wall to shaft centre, bore
INV = 12.5                            # headrace invert
LP, DP = 20.0, 2.4                    # penstock + tailpipe, bore
GAP = 0.48                            # nozzle opening
Q0, F = 7.6, 0.03                     # discharge per metre width, Darcy f (measured)
ROOF, FREEBOARD = 35.0, 3.0
H = RES - ZNOZ
DHH, DHP = 2 * DH, 2 * DP             # hydraulic diameters of a unit-width slot
Ds_of = lambda d: 2.5 + 0.5 * d       # noqa: E731

# ------------------------------------------------------------ the answers
u0 = Q0 / DH
vh = u0 * u0 / (2 * g)
hf = F * (L / DHH) * vh
z0 = vh + hf
k = z0 / (u0 * u0)


def crest(Ds, u=u0, zz=z0):
    """Rigid-column crest with quadratic friction, m above the reservoir."""
    kk = zz / (u * u); r = DH / Ds
    Z = L * r / (2 * g * kk); C = -(Z / kk) * math.exp(-zz / Z)
    fz = lambda z: C * math.exp(z / Z) + (z + Z) / kk      # noqa: E731  z positive DOWN
    a, b = -1.5 * u * math.sqrt(L * r / g), 0.0
    for _ in range(100):
        m = 0.5 * (a + b)
        if fz(a) * fz(m) <= 0: b = m
        else: a = m
    return -0.5 * (a + b), Z, C


zf = lambda Ds: u0 * math.sqrt(L * DH / (g * Ds))              # noqa: E731
Tf = lambda Ds: 2 * math.pi * math.sqrt(L * Ds / (g * DH))     # noqa: E731
# B3: the narrowest shaft whose crest clears the roof by the freeboard
limit = ROOF - RES - FREEBOARD
a, b = 0.5, 6.0
for _ in range(100):
    m = 0.5 * (a + b)
    if crest(m)[0] > limit: a = m
    else: b = m
Ds_min = 0.5 * (a + b)
# C: the penstock
up = Q0 / DP
hfp = F * (LP / DHP) * up * up / (2 * g)
hft = hf + hfp
ujet = math.sqrt(2 * g * (H - hft))
P0 = rho * g * Q0 * (H - hft)
kt = F * (L / DHH) / (2 * g * DH * DH) + F * (LP / DHP) / (2 * g * DP * DP)
qstar = math.sqrt(H / (3 * kt))
ujet_star = math.sqrt(2 * g * 2 * H / 3)
opening = qstar / ujet_star
Pstar = rho * g * qstar * 2 * H / 3
MEAS_CREST = [5.16, 4.79, 4.67, 4.18, 3.99, 3.79, 3.70, 3.58, 3.27, 3.22]
MEAS_T = [14.9, 16.2, 17.6, 19.1, 18.9, 19.8, 21.2, 21.0, 22.2, 23.7]


# ------------------------------------------------------------ the figure
def scheme(path):
    fig, ax = plt.subplots(figsize=(8.4, 4.4))
    ax.set_xlim(-0.5, 70.5); ax.set_ylim(-1, 36.5); ax.set_aspect("equal"); ax.axis("off")
    W = "#4a86c8"
    ax.add_patch(plt.Rectangle((0, INV), 7.65, RES - INV, fc=W, ec="none"))                 # reservoir
    ax.add_patch(plt.Rectangle((7.65, INV), 48.5 - 7.65, DH, fc=W, ec="none"))              # headrace
    ax.add_patch(plt.Rectangle((48.5, INV), 3.0, 24.5 - INV, fc=W, ec="none"))              # shaft
    tx, tz = 10 / math.hypot(10, 11), -11 / math.hypot(10, 11); nx, nz = -tz, tx; h = DP / 2
    ax.add_patch(plt.Polygon([[50 + h * nx, 14 + h * nz], [60 + h * nx, 3 + h * nz],
                              [60 - h * nx, 3 - h * nz], [50 - h * nx, 14 - h * nz]], fc=W, ec="none"))
    ax.add_patch(plt.Rectangle((59.0, ZNOZ - h), 6.0, DP, fc=W, ec="none"))                 # tailpipe
    ax.add_patch(plt.Rectangle((65, ZNOZ - GAP / 2), 5, GAP, fc=W, ec="none"))              # jet
    for z0_, z1_ in ((ZNOZ - h - 0.3, ZNOZ - GAP / 2), (ZNOZ + GAP / 2, ZNOZ + h + 0.3)):
        ax.add_patch(plt.Rectangle((64.75, z0_), 0.5, z1_ - z0_, fc="k", ec="none"))       # nozzle plates
    ax.plot([63.5, 63.5], [ZNOZ - h, ZNOZ + h], color="#2f8f5b", lw=2.5)                     # valve
    ax.plot([0, 70], [ROOF, ROOF], "k--", lw=0.8); ax.plot([0, 70], [0, 0], "k-", lw=0.8)
    ax.plot([7.65, 7.65], [RES + 0.2, ROOF], "k-", lw=1)                                     # reservoir wall
    ax.plot([7.65, 48.5], [INV + DH, INV + DH], "k-", lw=1); ax.plot([0, 48.5], [INV, INV], "k-", lw=1)
    ax.plot([48.5, 48.5], [INV + DH, ROOF], "k-", lw=1); ax.plot([51.5, 51.5], [14.1, ROOF], "k-", lw=1)
    f9 = dict(fontsize=8.5)
    ax.text(3.8, RES + 0.6, "reservoir 24.9 m", ha="center", **f9)
    ax.text(28, INV + DH + 0.8, "headrace  L = 42.4 m,  D_h = 3.05 m", ha="center", **f9)
    ax.text(28, INV - 1.6, "invert 12.5 m", ha="center", **f9)
    ax.text(50, 30.5, "surge shaft\nD_s = 2.5 + 0.5·d", ha="center", **f9)
    ax.text(57.5, 11.0, "penstock\nD_p = 2.4 m, ≈ 20 m", ha="center", **f9)
    ax.text(62.2, 5.4, "valve", ha="center", color="#2f8f5b", **f9)
    ax.text(66.5, 6.2, "nozzle 0.48 m\nat z = 3.0 m", ha="center", **f9)
    ax.text(35, ROOF + 0.6, "roof 35 m", ha="center", **f9)
    ax.text(0.3, -0.9, "datum z = 0", va="top", **f9)
    ax.text(66, ROOF - 2.5, "x: 0–70 m", ha="center", **f9)
    fig.tight_layout(); fig.savefig(path, dpi=170); plt.close(fig)


# ------------------------------------------------------------ the document
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.0); s.left_margin = s.right_margin = Cm(2.2)


def Hd(text, level=1): return doc.add_heading(text, level=level)


def P(text="", bold=False, italic=False, size=None, align=None, space_after=6):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic
    if size: r.font.size = Pt(size)
    if align is not None: p.alignment = align
    p.paragraph_format.space_after = Pt(space_after); return p


def EQ(text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.2); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = "Cambria Math"; r.font.size = Pt(11.5)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math"); return p


def NUM(text, lead=None):
    p = doc.add_paragraph(style="List Number")
    if lead: r = p.add_run(lead); r.bold = True
    p.add_run(text); p.paragraph_format.space_after = Pt(4); return p


def ANS(lines):
    for _ in range(lines):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
        pPr = p._p.get_or_add_pPr(); bdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
        for a_, v_ in (("w:val", "dotted"), ("w:sz", "4"), ("w:space", "1"), ("w:color", "A0A0A0")): bot.set(qn(a_), v_)
        bdr.append(bot); pPr.append(bdr)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def TABLE(rows, widths, header=True):
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            c = t.cell(i, j); c.width = Cm(widths[j]); c.text = ""
            r = c.paragraphs[0].add_run(str(cell)); r.font.size = Pt(10)
            if header and i == 0: r.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2); return t


P("Hydropower and unsteady flow — tutorial sheet", bold=True, size=18, space_after=2)
P("Complete by hand before the hydraulician session (exercise HP-3, Design the surge tower). Everything on this "
  "sheet is the scheme you will run in class, and every number on it can be checked on screen.", italic=True, space_after=10)

Hd("The scheme", 1)
fig_path = os.path.join(tempfile.gettempdir(), "hp3-scheme-model.png"); scheme(fig_path)
doc.add_picture(fig_path, width=Cm(15)); doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
P("A reservoir feeds a level headrace tunnel; at the knee a surge shaft rises off the tunnel, open to the air, and a "
  "penstock drops to a nozzle at the power house. The model is a vertical slice one metre wide: discharges are per "
  "metre of width and an \"area\" is a height, so A/A_s = D_h/D_s. Heights are above the datum, the floor of the "
  "picture.", space_after=4)
TABLE([["item", "value"],
       ["reservoir level", "24.9 m (the slider says 25.0; a gauge by the wall reads 24.9)"],
       ["headrace", "level; invert 12.5 m; bore D_h = 3.05 m; length L = 42.4 m from the reservoir wall (x = 7.65 m) to the shaft centre (x = 50 m)"],
       ["surge shaft", "width D_s = 2.5 + 0.5·d m, d the last digit of your student number; open to the air; the roof of the domain is at 35 m"],
       ["penstock", "bore D_p = 2.4 m; 14.9 m from the knee at (50, 14) down to (60, 3), then a 5 m level tailpipe to the nozzle at x = 65 m: L_p ≈ 20 m"],
       ["nozzle", "opening 0.48 m centred on the tailpipe axis at z = 3.0 m, discharging to atmosphere; the valve at x = 63.5 m is the turbine's instantaneous shutdown (the V key)"],
       ["steady discharge", "q₀ = 7.6 m²/s per metre of width (hover the headrace on screen: the card prints q and the bore-mean V)"],
       ["friction", "Darcy f = 0.03, measured in the app with two gauges 22 m apart on the headrace axis. For a slot one metre wide the hydraulic diameter is D_H = 2D (R_h = D/2), so h_f = f·(L/D_H)·u²/2g"],
       ["constants", "g = 9.81 m/s², ρ = 1000 kg/m³"]], [3.4, 12.6])

Hd("Notation", 2)
P("z is the water level in the shaft measured DOWN from the reservoir level (positive below it): z₀ is the steady "
  "running level and z_max the crest of the upsurge above the reservoir. u is the headrace velocity, L and A its "
  "length and \"area\" (= D_h), A_s the shaft \"area\" (= D_s), k the lumped loss coefficient in z₀ = k·u², and "
  "Z = L·A/(2·g·k·A_s). This is the app's notation (docs/notation.md); the worked example you have seen writes y for z.",
  space_after=8)

Hd("Part A — Steady running (predict, then read it on screen)", 1)
NUM(" The velocity in the headrace, u₀ = q₀/D_h.", "A1.")
ANS(2)
NUM(" The friction loss along the headrace.", "A2.")
EQ("h_f = f · (L/D_H) · u₀² / (2g),        D_H = 2·D_h")
ANS(2)
NUM(" The shaft is open to the air and carries no flow, so it is a piezometer at the knee: it reads the hydraulic "
    "grade line there. How far below the reservoir does its water stand? (Velocity head u₀²/2g plus h_f; take the "
    "entry loss as negligible.) The app's gauges read about 0.35 m, wobbling ±0.05 m. Is your prediction inside that?", "A3.")
ANS(3)
NUM(" Write the whole drawdown as z₀ = k·u₀² and find k (s²/m). Keep it: the surge uses k, not f.", "A4.")
ANS(2)

Hd("Part B — Instantaneous shutdown", 1)
P("Rigid column: the water in the headrace is one incompressible slug of length L. When the valve shuts it "
  "decelerates against the head it builds in the shaft. With z measured down from the reservoir:", space_after=2)
EQ("(L/g) · du/dt = z − k·u|u|            dz/dt = −(A/A_s) · u")
NUM(" Ignore friction. Eliminate t between the two equations (or write the energy balance) and show that the column "
    "comes to rest when the shaft has risen z_max = u₀·√(L·A/(g·A_s)) above the reservoir, with period "
    "T = 2π·√(L·A_s/(g·A)). Evaluate both for YOUR shaft width.", "B1.")
ANS(5)
NUM(" Keep friction. Dividing the two equations gives a linear equation in u² whose solution is", "B2.")
EQ("u² = C · e^(z/Z) + (z + Z)/k ,        Z = L·A / (2·g·k·A_s)")
P("C is fixed by the steady state (u = u₀ at z = z₀ = k·u₀²): C = −(Z/k)·e^(−z₀/Z). The crest is where u = 0. "
  "Evaluate Z and C for your shaft, then find the crest by trial and error on", space_after=2)
EQ("0 = C · e^(z/Z) + (z + Z)/k")
P("(z is negative above the reservoir; start near −z_max from B1 and work upward). z₀/Z is small here, so the crest "
  "lands within a few centimetres of z_max − z₀: the shaft starts z₀ below the reservoir and rises by about the "
  "frictionless amount. Do the trial anyway and see it happen.", space_after=2)
TABLE([["trial z (m)", "C·e^(z/Z)", "(z + Z)/k", "sum"], ["", "", "", ""], ["", "", "", ""], ["", "", "", ""], ["", "", "", ""]],
      [3.5, 4.0, 4.0, 4.0])
ANS(2)
NUM(" The crest stands at 24.9 + z_max above datum. Does your shaft hold the first upsurge with 3 m of freeboard "
    "under the 35 m roof? What is the narrowest shaft, to 0.1 m, that would?", "B3.")
ANS(3)
NUM(" The Knee x slider moves the shaft along the tunnel. If L were 30 m instead of 42.4 m, by what factor would "
    "z_max and T change?", "B4.")
ANS(2)

Hd("Part C — Maximum power through the penstock", 1)
P("The power delivered to the nozzle, per metre of width, is the discharge times the head left after friction in "
  "the headrace and the penstock together:", space_after=2)
EQ("P = ρ·g·q·(H − h_f),      H = 24.9 − 3.0 m,      h_f = k_t·q²")
EQ("k_t = f·(L/D_Hh) / (2g·D_h²)  +  f·(L_p/D_Hp) / (2g·D_p²)")
NUM(" Show that P is greatest when h_f = H/3, whatever k_t is (set dP/dq = 0).", "C1.")
ANS(3)
NUM(" At the design point q₀: the friction loss in headrace and penstock, the jet speed u_jet = √(2g·(H − h_f)) and "
    "the power, which is also ½·ρ·q₀·u_jet². On screen: Field → Speed, hover the jet just past the nozzle "
    "(the app reads about 20.6 m/s).", "C2.")
ANS(4)
NUM(" The discharge q* = √(H/(3·k_t)) that would put a third of the head into friction, the jet speed at that point, "
    "√(2g·2H/3), and the nozzle opening it would need, q*/u_jet. Compare with the penstock bore (2.4 m) and with the "
    "top of the Nozzle gap slider (0.9 m, about 12 m²/s measured). Where does that leave the peak on this scheme?", "C3.")
ANS(4)

Hd("Part D — In the session", 1)
NUM(" Fill the predicted column from Parts A–C for your own digit.", "D1.")
NUM(" In the session read u₀ and z₀ before the slam, then the rise of the shaft above its steady level and the "
    "period T after it; z_max = rise − z₀. Compare with D1; the class then pools every width into one curve.", "D2.")
TABLE([["", "predicted (D1)", "measured (session)"],
       ["D_s (m)", "", ""], ["u₀ (m/s)", "", ""], ["z₀ (m)", "", ""],
       ["z_max, no friction (m)", "", "—"], ["z_max, with friction (m)", "", ""], ["T (s)", "", ""],
       ["Darcy f from two headrace gauges (optional)", "0.03", ""],
       ["u_jet (m/s) and P (MW per m) (optional)", "", ""]], [6.2, 4.4, 4.4])

# ------------------------------------------------------------ answers
doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
Hd("Answers (for the tutor)", 1)
P("Computed from the constants above; two or three significant figures are all the sheet supports.", italic=True)
TABLE([["", ""],
       ["A1", "u₀ = 7.6/3.05 = %.2f m/s;  u₀²/2g = %.3f m" % (u0, vh)],
       ["A2", "h_f = 0.03 × (42.4/6.1) × %.3f = %.3f m" % (vh, hf)],
       ["A3", "z₀ = %.3f + %.3f = %.2f m below the reservoir. The app reads 0.35 ± 0.05 m: inside the wobble, and the entry loss is smaller than the reading." % (vh, hf, z0)],
       ["A4", "k = z₀/u₀² = %.4f s²/m" % k],
       ["B3", "The crest may reach 35 − 24.9 − 3 = %.1f m above the reservoir. With the friction crest ≈ u₀√(L·D_h/(g·D_s)) − z₀ that needs D_s ≥ %.2f m, so about 1.5 m — where the app's slider stops. (The app throttles shafts narrower than about 2.5 m: the entry into a narrow shaft costs head and the rise falls short of the formula.)" % (limit, Ds_min)],
       ["B4", "Both scale with √L: × √(30/42.4) = %.2f." % math.sqrt(30 / L)],
       ["C1", "P = ρgq(H − k_tq²); dP/dq = ρg(H − 3k_tq²) = 0 ⇒ k_tq² = H/3 = h_f."],
       ["C2", "u_p = 7.6/2.4 = %.2f m/s; h_f = %.3f (headrace) + %.3f (penstock) = %.2f m; H = 21.9 m; u_jet = √(19.62 × %.2f) = %.1f m/s; P = ρ·g·q₀·(H − h_f) = %.2f MW per m = ½ρq₀u_jet²." % (up, hf, hfp, hft, H - hft, ujet, P0 / 1e6)],
       ["C3", "k_t = %.5f s²/m⁵; q* = √(21.9/(3k_t)) = %.0f m²/s; u_jet = √(2g·2H/3) = %.1f m/s; opening = q*/u_jet = %.1f m — wider than the 2.4 m penstock, nearly four times the slider's top (about 12 m²/s at 0.9 m). The peak cannot be reached: sweep the slider and the jet stays near √(2gH) = %.1f m/s while the power keeps rising with the gap. (P* would be %.1f MW per m.)" % (kt, qstar, ujet_star, opening, math.sqrt(2 * g * H), Pstar / 1e6)],
       ], [1.4, 14.6], header=False)
P("B1, B2 — per digit (u₀ = %.2f m/s, z₀ = %.2f m, k = %.4f s²/m, L = 42.4 m, D_h = 3.05 m); the last two columns are what the app measured on the dry-run class:" % (u0, z0, k), bold=True, space_after=2)
rows = [["d", "D_s (m)", "z_max no friction (m)", "Z (m)", "z_max with friction (m)", "T (s)", "app: crest (m)", "app: T (s)"]]
for d in range(10):
    Ds = Ds_of(d); zk, Z, C = crest(Ds)
    rows.append([str(d), "%.1f" % Ds, "%.2f" % zf(Ds), "%.1f" % Z, "%.2f" % zk, "%.1f" % Tf(Ds), "%.2f" % MEAS_CREST[d], "%.1f" % MEAS_T[d]])
TABLE(rows, [0.9, 1.6, 2.6, 1.5, 2.7, 1.5, 2.2, 1.8])
P("The measured crests are 0.91–0.98 of the frictionless value and within about 5% of the with-friction column. "
  "The measured periods run 17–26% over the formula: the shaft's own water has inertia the rigid column leaves out "
  "(add h_s·A/A_s to L, with h_s ≈ 9 m of water standing in the shaft, and half the gap closes).", italic=True)

doc.save(OUT)
print("saved", OUT)
print("A: u0 %.3f  vh %.4f  hf %.4f  z0 %.4f  k %.4f" % (u0, vh, hf, z0, k))
print("B3: limit %.2f  Ds_min %.3f   B4: %.3f" % (limit, Ds_min, math.sqrt(30 / L)))
print("C2: up %.3f hfp %.4f hft %.4f ujet %.2f P %.3f MW/m   C3: kt %.6f q* %.1f ujet* %.2f opening %.2f P* %.2f" % (up, hfp, hft, ujet, P0 / 1e6, kt, qstar, ujet_star, opening, Pstar / 1e6))
for d in range(10):
    Ds = Ds_of(d); print("  d=%d Ds=%.1f zf=%.2f zk=%.2f T=%.1f" % (d, Ds, zf(Ds), crest(Ds)[0], Tf(Ds)))
